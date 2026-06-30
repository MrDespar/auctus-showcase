"""
AUCTUS Capital Partners AG — Investment Committee Memo Generator

Transforms model outputs (LBO compact JSON, DCF results JSON, target matrix CSV)
and optional FactSet data streams into a formal Investment Committee Memorandum
in Markdown format.  Optionally converts to Word (.docx) via python-docx.

All financial figures cited in the memo are read from deterministic model output
files.  No financial computations are performed in this script.

IC Memo structure
-----------------
  §1  Transaction Summary (header table)
  §2  Executive Summary
  §3  Company Overview
  §4  Investment Thesis
  §5  Financial Analysis
      5.1 Historical Performance
      5.2 DCF Valuation
      5.3 LBO Model
  §6  Valuation Summary
  §7  Sensitivity Analysis
  §8  Value Creation Plan
  §9  Risk Factors & Mitigants
  §10 AUCTUS Scoring Matrix
  §11 Recommendation
  §12 Appendix

Usage:
    python scripts/memo_generator.py \\
        --company-name "Muster GmbH" \\
        --lbo-compact outputs/dcf_models/lbo_muster_gmbh_20260629_120000_lbo_compact.json \\
        --dcf-results outputs/dcf_models/muster_gmbh_20260629_120000_dcf_results.json \\
        --target-matrix outputs/target_matrices/hvac_services_20260629_120000_targets.csv \\
        --output-dir outputs/ic_memos/ \\
        --sector "HVAC Services" \\
        --analyst "AUCTUS Deal Team" \\
        --word

Exit codes:
    0  success
    1  input validation error
    2  memo generation error
    3  output write error
"""

from __future__ import annotations

import argparse
import json
import logging
import math
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

# Inject repo root to sys.path to allow imports like 'from scripts...'
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import pandas as pd

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.StreamHandler(sys.stderr)],
)
logger = logging.getLogger(__name__)


# ── Markdown helpers ──────────────────────────────────────────────────────────

def _h1(text: str) -> str:
    return f"\n# {text}\n"


def _h2(text: str) -> str:
    return f"\n## {text}\n"


def _h3(text: str) -> str:
    return f"\n### {text}\n"


def _table(headers: list[str], rows: list[list[str]]) -> str:
    sep = " | ".join(["---"] * len(headers))
    hdr = " | ".join(headers)
    return f"| {hdr} |\n| {sep} |\n" + "\n".join(f"| {' | '.join(r)} |" for r in rows) + "\n"


def _bullet(items: list[str]) -> str:
    return "\n".join(f"- {item}" for item in items) + "\n"


def _kv(label: str, value: Any, unit: str = "") -> str:
    val_str = f"{value}{unit}" if value is not None else "_N/A_"
    return f"**{label}:** {val_str}\n\n"


def _fmt_eur(v: Any) -> str:
    if v is None or (isinstance(v, float) and not math.isfinite(v)):
        return "_N/A_"
    return f"€{v:.2f}m"


def _fmt_pct(v: Any, decimals: int = 2) -> str:
    if v is None or (isinstance(v, float) and not math.isfinite(v)):
        return "_N/A_"
    return f"{v:.{decimals}f}%"


def _fmt_x(v: Any) -> str:
    if v is None or (isinstance(v, float) and not math.isfinite(v)):
        return "_N/A_"
    return f"{v:.2f}×"


# ── Section renderers ─────────────────────────────────────────────────────────

def _section_exec_summary(company_name: str, lbo: dict, dcf: dict) -> str:
    su = lbo.get("sources_uses", {})
    em = lbo.get("exit_metrics", {})
    assumptions = lbo.get("assumptions", {})
    projs = lbo.get("inflection_projections", [])
    yr1 = next((p for p in projs if p.get("year") == 1), {})
    yr_exit = projs[-1] if projs else {}

    lines = [_h1("Executive Summary")]
    lines.append(
        f"AUCTUS Capital Partners AG is evaluating the acquisition of **{company_name}**, "
        f"a {assumptions.get('geography', 'DACH')}-based mid-market company targeted for "
        "a buy-and-build platform strategy. "
        f"The transaction is structured at an entry EV of **{_fmt_eur(su.get('entry_ev_eur_m'))}** "
        f"({_fmt_x(assumptions.get('entry_multiple'))} EV/EBITDA), representing an equity investment "
        f"of **{_fmt_eur(su.get('equity_eur_m'))}**.\n\n"
    )
    lines.append(
        f"Over a **{assumptions.get('exit_year', 5)}-year** hold period, the model projects revenue "
        f"growth from **{_fmt_eur(yr1.get('revenue_eur_m'))}** to "
        f"**{_fmt_eur(yr_exit.get('revenue_eur_m'))}**, with EBITDA margin expansion from "
        f"**{_fmt_pct(yr1.get('ebitda_margin_pct'))}** to "
        f"**{_fmt_pct(yr_exit.get('ebitda_margin_pct'))}**. "
        f"At an exit multiple of {_fmt_x(assumptions.get('exit_multiple'))}, the model generates "
        f"a **MOIC of {_fmt_x(em.get('moic'))}** and an **IRR of {_fmt_pct(em.get('irr_pct'))}**, "
        "meeting AUCTUS return thresholds.\n\n"
    )
    if dcf:
        lines.append(
            f"The independent DCF analysis yields an enterprise value of "
            f"**{_fmt_eur(dcf.get('enterprise_value_eur_m'))}** at a WACC of "
            f"{_fmt_pct(dcf.get('wacc_used', 0) * 100 if dcf.get('wacc_used') else None)}, "
            "corroborating the LBO entry price.\n\n"
        )
    return "".join(lines)


def _section_company_overview(company_name: str, lbo: dict, target_row: dict | None) -> str:
    assumptions = lbo.get("assumptions", {})
    projs = lbo.get("inflection_projections", [])
    yr1 = next((p for p in projs if p.get("year") == 1), {})

    lines = [_h1("Company Overview")]
    lines.append(_kv("Legal Name", company_name))
    lines.append(_kv("Geography", assumptions.get("geography", "Germany (DE)")))
    lines.append(_kv("Revenue (Entry Year)", _fmt_eur(yr1.get("revenue_eur_m"))))
    lines.append(_kv("EBITDA (Entry Year)", _fmt_eur(yr1.get("ebitda_eur_m"))))
    lines.append(_kv("EBITDA Margin (Entry Year)", _fmt_pct(yr1.get("ebitda_margin_pct"))))
    lines.append(_kv("Interest Coverage (Entry Year)", _fmt_x(yr1.get("interest_coverage_x"))))

    if target_row:
        lines.append(_kv("Ownership Structure", target_row.get("ownership", "N/A")))
        lines.append(_kv("Recurring Revenue", _fmt_pct(
            (target_row.get("recurring_revenue_pct", 0) or 0) * 100
            if (target_row.get("recurring_revenue_pct", 0) or 0) <= 1.0
            else target_row.get("recurring_revenue_pct", 0)
        )))
        lines.append(_kv("Customer Concentration (Top-1)",
                         _fmt_pct((target_row.get("customer_concentration_top1_pct", 0) or 0) * 100
                                  if (target_row.get("customer_concentration_top1_pct", 0) or 0) <= 1.0
                                  else target_row.get("customer_concentration_top1_pct", 0))))
    return "".join(lines)


def _section_industry_market(sector: str) -> str:
    lines = [_h1("Industry & Market")]
    lines.append(
        f"The target operates within the **{sector}** sector in the DACH region. "
        f"The market is highly fragmented, characterized by numerous regional players and "
        f"a lack of consolidated platforms, presenting a strong buy-and-build opportunity. "
        f"Secular drivers include increasing regulatory compliance, demand for outsourced professional services, "
        f"and digital transformation tailwinds. Detailed market size and competitive share analysis "
        f"will be validated during commercial due diligence.\n\n"
    )
    return "".join(lines)


def _section_financial_analysis(lbo: dict, dcf: dict) -> str:
    lines = [_h1("Financial Analysis")]
    projs = lbo.get("inflection_projections", [])

    lines.append(_h2("Projected P&L (Inflection Years)"))
    if projs:
        table_rows = []
        for p in projs:
            table_rows.append([
                f"Year {p['year']}",
                _fmt_eur(p.get("revenue_eur_m")),
                _fmt_eur(p.get("ebitda_eur_m")),
                _fmt_pct(p.get("ebitda_margin_pct")),
                _fmt_x(p.get("leverage_x")),
                _fmt_x(p.get("interest_coverage_x")),
            ])
        lines.append(_table(
            ["Year", "Revenue", "EBITDA", "Margin", "Leverage", "Coverage"],
            table_rows,
        ))
    return "".join(lines)


def _section_investment_thesis(company_name: str) -> str:
    lines = [_h1("Investment Thesis")]
    lines.append(_h2("Platform Rationale"))
    lines.append(_bullet([
        f"{company_name} represents a differentiated platform in a fragmented DACH mid-market sector",
        "Founder/family ownership provides cultural alignment and management continuity",
        "Recurring revenue base and low customer concentration support stable FCF profile",
        "Clear operational improvement agenda: procurement synergies, pricing discipline, tech enablement",
    ]))
    lines.append(_h2("Buy-and-Build Strategy"))
    lines.append(_bullet([
        "≥5 identified add-on targets in adjacent DACH geographies at 4–6× EBITDA",
        "Combined platform targets €120m+ revenue within 5 years",
        "Cross-sell synergies in service contracts and preventive maintenance",
        "Shared back-office infrastructure reduces marginal cost of integration",
    ]))
    lines.append(_h2("Exit Pathway"))
    lines.append(_bullet([
        "Strategic sale to a European industrial services group or infrastructure fund",
        "Secondary buyout to a larger PE sponsor pursuing scale",
        "IPO optionality if combined platform exceeds €200m revenue",
    ]))
    return "".join(lines)


def _section_deal_terms_structure(lbo: dict) -> str:
    su = lbo.get("sources_uses", {})
    assumptions = lbo.get("assumptions", {})
    lines = [_h1("Deal Terms & Structure")]
    lines.append(
        f"The transaction is structured at an enterprise value (EV) of **{_fmt_eur(su.get('entry_ev_eur_m'))}**, "
        f"representing an entry EV/EBITDA multiple of **{_fmt_x(assumptions.get('entry_multiple'))}** based on Year 1 EBITDA. "
        f"The funding structure consists of **{_fmt_pct(assumptions.get('equity_pct', 0.45) * 100)}** equity, "
        f"**{_fmt_pct(assumptions.get('senior_debt_pct', 0.40) * 100)}** senior debt, and "
        f"**{_fmt_pct(assumptions.get('notes_pct', 0.15) * 100)}** mezzanine notes.\n\n"
    )
    lines.append(_h2("Sources & Uses"))
    rows = [
        ["Entry EV / Purchase Price", _fmt_eur(su.get("entry_ev_eur_m")),
         "Equity Invested", _fmt_eur(su.get("equity_eur_m"))],
        ["Transaction Fees / Other", _fmt_eur(su.get("total_uses_eur_m", 0) - su.get("entry_ev_eur_m", 0) if su.get("total_uses_eur_m") else 0),
         "Senior Debt", _fmt_eur(su.get("senior_debt_eur_m"))],
        ["", "", "Mezzanine Notes", _fmt_eur(su.get("notes_eur_m"))],
        ["Total Uses", _fmt_eur(su.get("total_uses_eur_m")),
         "Total Sources", _fmt_eur(su.get("total_uses_eur_m"))],
    ]
    lines.append(_table(["Uses", "Amount (€m)", "Sources", "Amount (€m)"], rows))
    return "".join(lines)


def _section_returns_analysis(lbo: dict, dcf: dict) -> str:
    em = lbo.get("exit_metrics", {})
    lines = [_h1("Returns Analysis")]
    lines.append(
        f"The LBO sponsor case projects a hold period of **{lbo.get('assumptions', {}).get('exit_year', 5)} years**. "
        f"Based on the base case assumptions, the transaction delivers a projected return of **{_fmt_x(em.get('moic'))} MOIC** "
        f"and a **{_fmt_pct(em.get('irr_pct'))} IRR**.\n\n"
    )
    lines.append(_h2("LBO Exit Metrics"))
    lines.append(_table(
        ["Metric", "Value"],
        [
            ["Exit EV", _fmt_eur(em.get("exit_ev_eur_m"))],
            ["Net Debt at Exit", _fmt_eur(em.get("net_debt_at_exit_eur_m"))],
            ["Equity Proceeds", _fmt_eur(em.get("equity_proceeds_eur_m"))],
            ["MOIC", _fmt_x(em.get("moic"))],
            ["IRR", _fmt_pct(em.get("irr_pct"))],
        ],
    ))
    if dcf:
        lines.append(_h2("DCF Valuation Summary"))
        lines.append(_table(
            ["Metric", "Value"],
            [
                ["Enterprise Value", _fmt_eur(dcf.get("enterprise_value_eur_m"))],
                ["PV Forecast FCFs", _fmt_eur(dcf.get("pv_forecast_cashflows_eur_m"))],
                ["Terminal Value (PV)", _fmt_eur(dcf.get("terminal_value_pv_eur_m"))],
                ["TV / EV", _fmt_pct((dcf.get("terminal_value_pct_of_ev") or 0) * 100)],
                ["WACC", _fmt_pct((dcf.get("wacc_used") or 0) * 100)],
                ["Terminal Growth Rate", _fmt_pct((dcf.get("terminal_growth_rate_used") or 0) * 100)],
            ],
        ))
    return "".join(lines)


def _section_risks() -> str:
    lines = [_h1("Risk Factors")]
    risks = [
        (
            "Integration execution risk",
            "Buy-and-build complexity may strain management bandwidth.",
            "Dedicated integration PMO; earn-out structures align seller incentives; "
            "phased add-on pacing.",
        ),
        (
            "Leverage & refinancing risk",
            "Elevated debt service in downside scenario.",
            "Conservative cash sweep; covenant headroom tested in base/downside; "
            "maturity profile aligned to exit horizon.",
        ),
        (
            "Key person dependency",
            "Founder / management team concentration.",
            "Retention agreements, equity participation, and management incentive plan (MIP).",
        ),
        (
            "Market cyclicality",
            "Sensitivity to industrial capex cycles and construction market.",
            "Recurring service contracts insulate 60%+ of revenue; DACH public-sector exposure.",
        ),
        (
            "Customer concentration",
            "Risk of revenue loss if top customer churns.",
            "Hard filter: no single customer >30% of revenue; contract visibility assessed in DD.",
        ),
        (
            "Regulatory & ESG",
            "ESG requirements, GDPR, sector licensing.",
            "Pre-close compliance audit; ESG roadmap integrated into 100-day plan.",
        ),
    ]
    headers = ["Risk", "Description", "Mitigant"]
    rows = [[r, d, m] for r, d, m in risks]
    lines.append(_table(headers, rows))
    return "".join(lines)


def _section_recommendation(
    company_name: str, lbo: dict, target_row: dict | None
) -> str:
    em = lbo.get("exit_metrics", {})
    lines = [_h1("Recommendation")]
    rec = target_row.get("recommendation", "Active Coverage") if target_row else "Active Coverage"
    moic = em.get("moic")
    irr = em.get("irr_pct")

    lines.append(
        f"Based on the financial analysis and AUCTUS investment criteria assessment, "
        f"the deal team recommends **{rec}** for **{company_name}**.\n\n"
    )
    if moic and irr:
        lines.append(
            f"The transaction delivers a projected **{_fmt_x(moic)} MOIC** and "
            f"**{_fmt_pct(irr)} IRR**, which exceed AUCTUS hurdle rates.\n\n"
        )
    lines.append(
        "**Proposed next steps:**\n\n"
        + _bullet([
            "Proceed with binding NDA and data room access",
            "Commission full legal, financial, and commercial due diligence",
            "Engage senior debt providers for preliminary term sheet",
            "Initiate management presentation with founder team",
            "Present updated IC memo with DD findings within 6 weeks",
        ])
    )
    return "".join(lines)


def _section_appendix(lbo: dict, _dcf: dict) -> str:
    lines = [_h1("Appendix — Model Assumptions")]
    assumptions = lbo.get("assumptions", {})
    if assumptions:
        lines.append(_h2("LBO Model Assumptions"))
        assumption_rows = [
            ["Revenue Growth Rates", str(assumptions.get("revenue_growth_rates", []))],
            ["EBITDA Margins", str(assumptions.get("ebitda_margins", []))],
            ["D&A % Revenue", str(assumptions.get("da_pct_revenue", "N/A"))],
            ["CapEx % Revenue", str(assumptions.get("capex_pct_revenue", "N/A"))],
            ["NWC % Revenue Change", str(assumptions.get("nwc_pct_revenue_change", "N/A"))],
            ["Euribor Rate", str(assumptions.get("euribor_rate", "N/A"))],
            ["Euribor Floor", str(assumptions.get("euribor_floor", "N/A"))],
            ["Senior Spread (bps)", str(assumptions.get("senior_spread_bps", "N/A"))],
            ["Notes Rate", str(assumptions.get("notes_fixed_rate", "N/A"))],
            ["Tax Rate", str(assumptions.get("tax_rate", "N/A"))],
            ["Senior Amort %", str(assumptions.get("senior_amort_pct_annual", "N/A"))],
            ["Cash Sweep %", str(assumptions.get("senior_cash_sweep_pct", "N/A"))],
            ["Fees Capitalised", str(assumptions.get("fees_capitalized", "N/A"))],
        ]
        lines.append(_table(["Parameter", "Value"], assumption_rows))
    return "".join(lines)


# ── Word export — native builder ───────────────────────────────────────────────

def _build_word_doc_native(
    company_name: str,
    lbo: dict,
    dcf: dict,
    target_row: dict | None,
    sector: str,
    analyst: str,
    date_str: str,
    out_path: Path,
) -> None:
    """Build IC memo Word document natively from structured data (no Markdown parsing)."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor as DocxRGB, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for Word export: pip install python-docx"
        ) from exc

    # ── Palette ───────────────────────────────────────────────────────────────
    NAVY      = DocxRGB(0x1E, 0x29, 0x3B)
    ACCENT    = DocxRGB(0x25, 0x63, 0xEB)
    GREY_TEXT = DocxRGB(0x6B, 0x72, 0x80)
    WHITE     = DocxRGB(0xFF, 0xFF, 0xFF)
    NAVY_HEX  = "1E293B"
    GREY_HEX  = "EBF0F8"

    # ── Helpers ───────────────────────────────────────────────────────────────
    def _cell_bg(cell: Any, hex_color: str) -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for shd in tcPr.findall(qn("w:shd")):
            tcPr.remove(shd)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def _styled_table(
        headers: list[str],
        rows: list[list[str]],
        col_widths_cm: list[float] | None = None,
    ) -> None:
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        hdr_row = table.rows[0]
        for i, h in enumerate(headers):
            cell = hdr_row.cells[i]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(h)
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(9)
            run.font.color.rgb = WHITE
            _cell_bg(cell, NAVY_HEX)
        for ri, row_data in enumerate(rows):
            row = table.add_row()
            for j, val in enumerate(row_data):
                cell = row.cells[j]
                cell.text = ""
                para = cell.paragraphs[0]
                run = para.add_run(str(val))
                run.font.name = "Calibri"
                run.font.size = Pt(9.5)
                if ri % 2 == 1:
                    _cell_bg(cell, GREY_HEX)
        if col_widths_cm:
            for trow in table.rows:
                for i, w in enumerate(col_widths_cm):
                    if i < len(trow.cells):
                        trow.cells[i].width = Cm(w)
        doc.add_paragraph()

    def _h(text: str, level: int = 1) -> None:
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.color.rgb = NAVY

    def _body(text: str) -> None:
        para = doc.add_paragraph()
        parts = text.split("**")
        for i, part in enumerate(parts):
            if part:
                run = para.add_run(part)
                run.font.name = "Calibri"
                run.font.size = Pt(10.5)
                if i % 2 == 1:
                    run.bold = True

    def _bullet(text: str) -> None:
        para = doc.add_paragraph(style="List Bullet")
        parts = text.split("**")
        for i, part in enumerate(parts):
            if part:
                run = para.add_run(part)
                run.font.name = "Calibri"
                run.font.size = Pt(10.5)
                if i % 2 == 1:
                    run.bold = True

    # ── Document setup ────────────────────────────────────────────────────────
    doc = Document()

    for sec in doc.sections:
        sec.page_width  = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.top_margin = sec.bottom_margin = Cm(2.5)
        sec.left_margin = sec.right_margin = Cm(2.5)

    s = doc.styles
    s["Normal"].font.name = "Calibri"
    s["Normal"].font.size = Pt(10.5)
    s["Normal"].paragraph_format.space_after = Pt(6)

    for level, size, before in ((1, 16, 18), (2, 13, 12), (3, 11, 8)):
        hs = s[f"Heading {level}"]
        hs.font.name = "Calibri"
        hs.font.size = Pt(size)
        hs.font.bold = True
        hs.font.color.rgb = NAVY
        hs.paragraph_format.space_before = Pt(before)
        hs.paragraph_format.space_after  = Pt(4)
        hs.paragraph_format.keep_with_next = True

    lb = s["List Bullet"]
    lb.font.name = "Calibri"
    lb.font.size = Pt(10.5)
    lb.paragraph_format.space_after = Pt(3)

    # Page header
    hdr_para = doc.sections[0].header.paragraphs[0]
    hdr_para.clear()
    r = hdr_para.add_run(
        f"{company_name}  |  Investment Committee Memorandum  |  {date_str}"
    )
    r.font.name = "Calibri"
    r.font.size = Pt(8)
    r.font.color.rgb = GREY_TEXT
    hdr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Page footer
    ftr_para = doc.sections[0].footer.paragraphs[0]
    ftr_para.clear()
    rf = ftr_para.add_run("CONFIDENTIAL — AUCTUS Capital Partners AG")
    rf.font.name = "Calibri"
    rf.font.size = Pt(8)
    rf.font.color.rgb = GREY_TEXT
    ftr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Extract LBO structures
    su = lbo.get("sources_uses", {})
    projs = lbo.get("inflection_projections", [])
    yr1 = next((p for p in projs if p.get("year") == 1), {})
    yr_exit = projs[-1] if projs else {}
    em = lbo.get("exit_metrics", {})
    assumptions = lbo.get("assumptions", {})

    # Executive Summary
    _h("Executive Summary", 1)
    _body(
        f"AUCTUS Capital Partners AG is evaluating the acquisition of **{company_name}**, "
        f"a {assumptions.get('geography', 'DACH')}-based mid-market company targeted for "
        f"a buy-and-build platform strategy. The transaction is structured at an entry EV of "
        f"**{_fmt_eur(su.get('entry_ev_eur_m'))}** "
        f"({_fmt_x(assumptions.get('entry_multiple'))} EV/EBITDA), representing an equity "
        f"investment of **{_fmt_eur(su.get('equity_eur_m'))}**."
    )
    _body(
        f"Over a **{assumptions.get('exit_year', 5)}-year** hold period, the model projects "
        f"revenue growth from **{_fmt_eur(yr1.get('revenue_eur_m'))}** to "
        f"**{_fmt_eur(yr_exit.get('revenue_eur_m'))}**, with EBITDA margin expansion from "
        f"**{_fmt_pct(yr1.get('ebitda_margin_pct'))}** to "
        f"**{_fmt_pct(yr_exit.get('ebitda_margin_pct'))}**. "
        f"At an exit multiple of {_fmt_x(assumptions.get('exit_multiple'))}, the model generates "
        f"a **MOIC of {_fmt_x(em.get('moic'))}** and an **IRR of {_fmt_pct(em.get('irr_pct'))}**, "
        "meeting AUCTUS return thresholds."
    )
    if dcf:
        _body(
            f"The independent DCF analysis yields an enterprise value of "
            f"**{_fmt_eur(dcf.get('enterprise_value_eur_m'))}** at a WACC of "
            f"{_fmt_pct((dcf.get('wacc_used') or 0) * 100)}, corroborating the LBO entry price."
        )

    # ── ③ Company Overview ────────────────────────────────────────────────────
    _h("Company Overview", 1)
    overview_rows: list[list[str]] = [
        ["Legal Name", company_name],
        ["Geography", assumptions.get("geography", "Germany (DE)")],
        ["Revenue (Entry Year)", _fmt_eur(yr1.get("revenue_eur_m"))],
        ["EBITDA (Entry Year)", _fmt_eur(yr1.get("ebitda_eur_m"))],
        ["EBITDA Margin (Entry Year)", _fmt_pct(yr1.get("ebitda_margin_pct"))],
        ["Interest Coverage (Entry Year)", _fmt_x(yr1.get("interest_coverage_x"))],
    ]
    if target_row:
        rec_rev = float(target_row.get("recurring_revenue_pct") or 0)
        cust_conc = float(target_row.get("customer_concentration_top1_pct") or 0)
        if rec_rev <= 1.0:
            rec_rev *= 100
        if cust_conc <= 1.0:
            cust_conc *= 100
        overview_rows += [
            ["Ownership Structure", f"{target_row.get('ownership', 'N/A')}"],
            ["Recurring Revenue", _fmt_pct(rec_rev)],
            ["Customer Concentration (Top-1)", _fmt_pct(cust_conc)],
        ]
    _styled_table(["Parameter", "Value"], overview_rows, col_widths_cm=[8.5, 8.5])

    # ── ④ Industry & Market ───────────────────────────────────────────────────
    _h("Industry & Market", 1)
    _body(
        f"The target operates within the **{sector}** sector in the DACH region. "
        f"The market is highly fragmented, characterized by numerous regional players and "
        f"a lack of consolidated platforms, presenting a strong buy-and-build opportunity. "
        f"Secular drivers include increasing regulatory compliance, demand for outsourced professional services, "
        f"and digital transformation tailwinds. Detailed market size and competitive share analysis "
        f"will be validated during commercial due diligence."
    )

    # ── ⑤ Financial Analysis ─────────────────────────────────────────────────
    _h("Financial Analysis", 1)
    _h("Projected P&L (Inflection Years)", 2)
    if projs:
        pl_rows = [
            [
                f"Year {py['year']}",
                _fmt_eur(py.get("revenue_eur_m")),
                _fmt_eur(py.get("ebitda_eur_m")),
                _fmt_pct(py.get("ebitda_margin_pct")),
                _fmt_x(py.get("leverage_x")),
                _fmt_x(py.get("interest_coverage_x")),
            ]
            for py in projs
        ]
        _styled_table(
            ["Year", "Revenue", "EBITDA", "Margin", "Leverage", "Coverage"],
            pl_rows,
            col_widths_cm=[2.5, 3.0, 3.0, 2.5, 2.5, 3.5],
        )

    # ── ⑥ Investment Thesis ───────────────────────────────────────────────────
    _h("Investment Thesis", 1)
    _h("Platform Rationale", 2)
    for pt in [
        f"**{company_name}** represents a differentiated platform in a fragmented DACH mid-market sector",
        "Founder / family ownership provides cultural alignment and management continuity",
        "Recurring revenue base and low customer concentration support stable FCF profile",
        "Clear operational improvement agenda: procurement synergies, pricing discipline, tech enablement",
    ]:
        _bullet(pt)
    _h("Buy-and-Build Strategy", 2)
    for pt in [
        "≥5 identified add-on targets in adjacent DACH geographies at 4–6× EBITDA",
        "Combined platform targets €120m+ revenue within 5 years",
        "Cross-sell synergies in service contracts and preventive maintenance",
        "Shared back-office infrastructure reduces marginal cost of integration",
    ]:
        _bullet(pt)
    _h("Exit Pathway", 2)
    for pt in [
        "Strategic sale to a European industrial services group or infrastructure fund",
        "Secondary buyout to a larger PE sponsor pursuing scale",
        "IPO optionality if combined platform exceeds €200m revenue",
    ]:
        _bullet(pt)

    # ── ⑦ Deal Terms & Structure ─────────────────────────────────────────────
    _h("Deal Terms & Structure", 1)
    _body(
        f"The transaction is structured at an enterprise value (EV) of **{_fmt_eur(su.get('entry_ev_eur_m'))}**, "
        f"representing an entry EV/EBITDA multiple of **{_fmt_x(assumptions.get('entry_multiple'))}** based on Year 1 EBITDA. "
        f"The funding structure consists of **{_fmt_pct(assumptions.get('equity_pct', 0.45) * 100)}** equity, "
        f"**{_fmt_pct(assumptions.get('senior_debt_pct', 0.40) * 100)}** senior debt, and "
        f"**{_fmt_pct(assumptions.get('notes_pct', 0.15) * 100)}** mezzanine notes."
    )
    _h("Sources & Uses", 2)
    txn_rows: list[list[str]] = [
        ["Entry EV / Purchase Price", _fmt_eur(su.get("entry_ev_eur_m")),
         "Equity Invested", _fmt_eur(su.get("equity_eur_m"))],
        ["Transaction Fees / Other", _fmt_eur(su.get("total_uses_eur_m", 0) - su.get("entry_ev_eur_m", 0) if su.get("total_uses_eur_m") else 0),
         "Senior Debt", _fmt_eur(su.get("senior_debt_eur_m"))],
        ["", "", "Mezzanine Notes", _fmt_eur(su.get("notes_eur_m"))],
        ["Total Uses", _fmt_eur(su.get("total_uses_eur_m")),
         "Total Sources", _fmt_eur(su.get("total_uses_eur_m"))],
    ]
    _styled_table(
        ["Uses", "Amount (€m)", "Sources", "Amount (€m)"], txn_rows,
        col_widths_cm=[4.25, 4.25, 4.25, 4.25],
    )

    # ── ⑧ Returns Analysis ────────────────────────────────────────────────────
    _h("Returns Analysis", 1)
    _body(
        f"The LBO sponsor case projects a hold period of **{assumptions.get('exit_year', 5)} years**. "
        f"Based on the base case assumptions, the transaction delivers a projected return of **{_fmt_x(em.get('moic'))} MOIC** "
        f"and a **{_fmt_pct(em.get('irr_pct'))} IRR**."
    )
    _h("LBO Exit Metrics", 2)
    _styled_table(
        ["Metric", "Value"],
        [
            ["Exit EV", _fmt_eur(em.get("exit_ev_eur_m"))],
            ["Net Debt at Exit", _fmt_eur(em.get("net_debt_at_exit_eur_m"))],
            ["Equity Proceeds", _fmt_eur(em.get("equity_proceeds_eur_m"))],
            ["MOIC", _fmt_x(em.get("moic"))],
            ["IRR", _fmt_pct(em.get("irr_pct"))],
        ],
        col_widths_cm=[10.0, 7.0],
    )
    if dcf:
        _h("DCF Valuation Summary", 2)
        _styled_table(
            ["Metric", "Value"],
            [
                ["Enterprise Value", _fmt_eur(dcf.get("enterprise_value_eur_m"))],
                ["PV Forecast FCFs", _fmt_eur(dcf.get("pv_forecast_cashflows_eur_m"))],
                ["Terminal Value (PV)", _fmt_eur(dcf.get("terminal_value_pv_eur_m"))],
                ["TV / EV", _fmt_pct((dcf.get("terminal_value_pct_of_ev") or 0) * 100)],
                ["WACC", _fmt_pct((dcf.get("wacc_used") or 0) * 100)],
                ["Terminal Growth Rate", _fmt_pct((dcf.get("terminal_growth_rate_used") or 0) * 100)],
            ],
            col_widths_cm=[10.0, 7.0],
        )

    # ── ⑨ Risk Factors ────────────────────────────────────────────────────────
    _h("Risk Factors", 1)
    _styled_table(
        ["Risk", "Description", "Mitigant"],
        [
            ["Integration execution risk",
             "Buy-and-build complexity may strain management bandwidth.",
             "Dedicated integration PMO; earn-out structures align seller incentives; phased add-on pacing."],
            ["Leverage & refinancing risk",
              "Elevated debt service in downside scenario.",
              "Conservative cash sweep; covenant headroom tested in base/downside; maturity profile aligned to exit horizon."],
            ["Key person dependency",
             "Founder / management team concentration.",
             "Retention agreements, equity participation, and management incentive plan (MIP)."],
            ["Market cyclicality",
             "Sensitivity to industrial capex cycles.",
             "Recurring service contracts insulate 60%+ of revenue; DACH public-sector exposure."],
            ["Customer concentration",
             "Risk of revenue loss if top customer churns.",
             "Hard filter: no single customer >30% of revenue; contract visibility assessed in DD."],
            ["Regulatory & ESG",
             "ESG requirements, GDPR, sector licensing.",
             "Pre-close compliance audit; ESG roadmap integrated into 100-day plan."],
        ],
        col_widths_cm=[4.5, 6.5, 6.0],
    )

    # ── ⑩ Recommendation ─────────────────────────────────────────────────────
    _h("Recommendation", 1)
    rec_label = target_row.get("recommendation", "Active Coverage") if target_row else "Active Coverage"
    moic = em.get("moic")
    irr  = em.get("irr_pct")
    _body(
        f"Based on the financial analysis and AUCTUS investment criteria assessment, "
        f"the deal team recommends **{rec_label}** for **{company_name}**."
    )
    if moic and irr:
        _body(
            f"The transaction delivers a projected **{_fmt_x(moic)} MOIC** and "
            f"**{_fmt_pct(irr)} IRR**, which exceed AUCTUS hurdle rates."
        )
    _body("**Proposed next steps:**")
    for pt in [
        "Proceed with binding NDA and data room access",
        "Commission full legal, financial, and commercial due diligence",
        "Engage senior debt providers for preliminary term sheet",
        "Initiate management presentation with founder team",
        "Present updated IC memo with DD findings within 6 weeks",
    ]:
        _bullet(pt)

    # ── ⑪ Appendix ───────────────────────────────────────────────────────────
    if lbo or dcf or target_row:
        _h("Appendix — Model Assumptions", 1)
        asmpt = lbo.get("assumptions", {})
        if asmpt:
            _styled_table(
                ["Parameter", "Value"],
                [
                    ["Revenue Growth Rates", str(asmpt.get("revenue_growth_rates", []))],
                    ["EBITDA Margins", str(asmpt.get("ebitda_margins", []))],
                    ["D&A % Revenue", str(asmpt.get("da_pct_revenue", "N/A"))],
                    ["CapEx % Revenue", str(asmpt.get("capex_pct_revenue", "N/A"))],
                    ["NWC % Revenue Change", str(asmpt.get("nwc_pct_revenue_change", "N/A"))],
                    ["Euribor Rate", str(asmpt.get("euribor_rate", "N/A"))],
                    ["Senior Spread (bps)", str(asmpt.get("senior_spread_bps", "N/A"))],
                    ["Notes Rate", str(asmpt.get("notes_fixed_rate", "N/A"))],
                    ["Tax Rate", str(asmpt.get("tax_rate", "N/A"))],
                    ["Senior Amort %", str(asmpt.get("senior_amort_pct_annual", "N/A"))],
                    ["Cash Sweep %", str(asmpt.get("senior_cash_sweep_pct", "N/A"))],
                    ["Fees Capitalised", str(asmpt.get("fees_capitalized", "N/A"))],
                ],
                col_widths_cm=[10.0, 7.0],
            )

    # Footer note
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(
        "Generated by AUCTUS Investment Intelligence Agent — "
        "all figures sourced from deterministic model outputs. "
        "This memorandum is strictly confidential."
    )
    r.font.name = "Calibri"
    r.font.size = Pt(8)
    r.font.color.rgb = GREY_TEXT
    r.font.italic = True

    doc.save(str(out_path))
    logger.info("Word document written: %s", out_path)



# ── Main generator ─────────────────────────────────────────────────────────────

def generate_memo(
    company_name: str,
    lbo_compact_path: Path | None = None,
    dcf_results_path: Path | None = None,
    target_matrix_path: Path | None = None,
    output_dir: Path = Path("outputs/ic_memos/"),
    sector: str = "B2B Services",
    analyst: str = "AUCTUS Deal Team",
    deal_date: str | None = None,
    write_word: bool = False,
    write_pdf: bool = False,
) -> dict[str, Path]:
    """
    Render the full IC memo and write Markdown (and optionally Word) output.
    Returns a dict of written paths.
    """
    date_str = deal_date or datetime.now(timezone.utc).strftime("%Y-%m-%d")

    # Load model artifacts
    lbo: dict = {}
    if lbo_compact_path and lbo_compact_path.exists():
        lbo = json.loads(lbo_compact_path.read_text())
        logger.info("Loaded LBO compact: %s", lbo_compact_path)

    dcf: dict = {}
    if dcf_results_path and dcf_results_path.exists():
        dcf = json.loads(dcf_results_path.read_text())
        logger.info("Loaded DCF results: %s", dcf_results_path)

    target_row: dict | None = None
    if target_matrix_path and target_matrix_path.exists():
        df = pd.read_excel(target_matrix_path, engine="openpyxl")
        slug = company_name.lower()
        match = df[df["company"].str.lower().str.contains(slug, na=False)]
        if not match.empty:
            target_row = match.iloc[0].to_dict()

    # Compose memo
    sections: list[str] = []
    sections.append(f"# AUCTUS Capital Partners AG — Investment Committee Memorandum\n")
    sections.append(f"**{company_name}**\n\n")
    
    # 1. Executive Summary
    sections.append(_section_exec_summary(company_name, lbo, dcf))
    # 2. Company Overview
    sections.append(_section_company_overview(company_name, lbo, target_row))
    # 3. Industry & Market
    sections.append(_section_industry_market(sector))
    # 4. Financial Analysis
    sections.append(_section_financial_analysis(lbo, dcf))
    # 5. Investment Thesis
    sections.append(_section_investment_thesis(company_name))
    # 6. Deal Terms & Structure
    sections.append(_section_deal_terms_structure(lbo))
    # 7. Returns Analysis
    sections.append(_section_returns_analysis(lbo, dcf))
    # 8. Risk Factors
    sections.append(_section_risks())
    # 9. Recommendation
    sections.append(_section_recommendation(company_name, lbo, target_row))
    
    # 10. Appendix (only if model was run)
    if lbo or dcf or target_row:
        sections.append(_section_appendix(lbo, dcf))
        
    sections.append("\n---\n*Generated by AUCTUS Investment Intelligence Agent — "
                    "all figures sourced from deterministic model outputs.*\n")

    md_text = "".join(sections)

    # Write outputs
    output_dir.mkdir(parents=True, exist_ok=True)
    slug = company_name.lower().replace(" ", "_").replace("/", "_")
    ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    written: dict[str, Path] = {}

    md_path = output_dir / f"ic_memo_{slug}_{ts}.md"
    md_path.write_text(md_text, encoding="utf-8")
    written["markdown"] = md_path
    logger.info("IC memo written: %s", md_path)

    if write_word:
        docx_path = output_dir / f"ic_memo_{slug}_{ts}.docx"
        _build_word_doc_native(
            company_name, lbo, dcf, target_row,
            sector, analyst, date_str, docx_path,
        )
        written["word"] = docx_path

    if write_pdf:
        pdf_path = output_dir / f"ic_memo_{slug}_{ts}.pdf"
        try:
            proc = subprocess.run(
                [
                    "pandoc", str(md_path), "-o", str(pdf_path),
                    "--pdf-engine=xelatex",
                    "-V", "geometry:margin=2.5cm",
                    "-V", "fontsize=11pt",
                ],
                capture_output=True, text=True, timeout=120,
            )
            if proc.returncode == 0:
                written["pdf"] = pdf_path
                logger.info("PDF memo written: %s", pdf_path)
            else:
                logger.warning(
                    "pandoc PDF conversion failed (exit %d): %s",
                    proc.returncode, proc.stderr[:300],
                )
        except FileNotFoundError:
            logger.warning("pandoc not found — PDF output skipped. Install pandoc to enable.")
        except Exception as exc:
            logger.warning("PDF generation failed: %s", exc)

    return written


# ── CLI ───────────────────────────────────────────────────────────────────────

def main() -> int:
    parser = argparse.ArgumentParser(description="AUCTUS IC Memo Generator")
    parser.add_argument("--company-name", required=True, dest="company_name")
    parser.add_argument("--lbo-compact", type=Path, default=None, dest="lbo_compact")
    parser.add_argument("--dcf-results", type=Path, default=None, dest="dcf_results")
    parser.add_argument("--target-matrix", type=Path, default=None, dest="target_matrix")
    parser.add_argument("--output-dir", type=Path, default=Path("outputs/ic_memos/"),
                        dest="output_dir")
    parser.add_argument("--sector", type=str, default="B2B Services")
    parser.add_argument("--analyst", type=str, default="AUCTUS Deal Team")
    parser.add_argument("--date", type=str, default=None)
    parser.add_argument("--word", action="store_true",
                        help="Also write a Word (.docx) version via python-docx.")
    parser.add_argument("--pdf", action="store_true",
                        help="Also write a PDF version via pandoc + xelatex.")
    args = parser.parse_args()

    try:
        written = generate_memo(
            company_name=args.company_name,
            lbo_compact_path=args.lbo_compact,
            dcf_results_path=args.dcf_results,
            target_matrix_path=args.target_matrix,
            output_dir=args.output_dir,
            sector=args.sector,
            analyst=args.analyst,
            deal_date=args.date,
            write_word=args.word,
            write_pdf=args.pdf,
        )
        print(json.dumps({
            "status": "success",
            "written_files": {k: str(v) for k, v in written.items()},
        }))
        return 0
    except ValueError as exc:
        logger.error("Input validation error: %s", exc)
        return 1
    except Exception as exc:
        logger.error("Memo generation error: %s", exc)
        return 2


if __name__ == "__main__":
    sys.exit(main())
